from model.year import Year
from model.lesson import Lesson
import regex as re
import simplejson as json
from docx import Document
import os


class FileParser:

    def __init__(self):
        pass

    def parseLessons(self, document, year):
        lection_num = None
        lessons = []
        for index, paragraph in enumerate(document.paragraphs):
            m = re.search(r'(\d+)\. +lekce[\t ]*(\w.*)', paragraph.text)
            if m and len(paragraph.text) < 150:
                if lection_num:
                    paragraphs = [paragraph.text for paragraph in document.paragraphs[curr_start:index-1]]
                    lesson = Lesson(year, lection_num, lection_name)
                    lesson.setContentFromParagraphs(paragraphs)
                    lessons.append(lesson)

                lection_num = int(m.group(1))
                lection_name = m.group(2)
                curr_start = index
            elif re.match(r'##', paragraph.text):
                paragraphs = [paragraph.text for paragraph in document.paragraphs[curr_start:index-1]]
                lesson = Lesson(year, lection_num, lection_name)
                lesson.setContentFromParagraphs(paragraphs)
                
                lessons.append(lesson)
        return lessons

    def parseDocument(self, year, filename):
        document = Document(filename)

        parsed = self.parseLessons(document, year)

        return Year(year, parsed)

    def parseDocumentGetYearFromFilename(self, filename):
        file = os.path.basename(filename)
        year = int(os.path.splitext(file)[0])

        return self.parseDocument(year, filename)
        
    