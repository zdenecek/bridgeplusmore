from docx import Document
import sqlite3
from classes import Database

print('Starting script...')

file = 'txt/Merrimac.docx'

document = Document(file)
db = Database('boards.sqlite')

for para in document.paragraphs:
    print(para.text)

for char in document.paragraphs[2].text:
    print(str(ord(char)) + f' - [{char}]')

db.close()

